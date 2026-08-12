"""
Universal CV Generator
======================
Reads a JSON data file and generates a formatted DOCX CV using the
AININ_SOFEA_CV.docx template for Word styles/themes.

Usage:
    python generate_cv.py --data <json_file> [--output <filename.docx>]
    python generate_cv.py --data cv_data.json
    python generate_cv.py --data cv_data.json --output ilham_effendy_cv.docx

Dependencies:
    pip install python-docx
"""

import argparse
import json
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("ERROR: python-docx not installed. Run: python -m pip install python-docx")
    sys.exit(1)

# ── Paths ─────────────────────────────────────────────────────────────
SCRIPT_DIR = Path(__file__).resolve().parent
WORKSPACE  = SCRIPT_DIR.parent
TEMPLATE   = WORKSPACE / "templates" / "AININ_SOFEA_CV.docx"
OUTPUT_DIR = WORKSPACE / "output"

# ── Constants ─────────────────────────────────────────────────────────
FONT = "Calibri"
FULL_WIDTH     = 10485
TWO_COL_LEFT   = 4508
TWO_COL_RIGHT  = 5977
REPORT_LEFT    = 7225
REPORT_RIGHT   = 3260
# (Removed ACT_COL constants since we use REPORT widths for all generic 2-col tables)


# ═══════════════════════════════════════════════════════════════════════
# LOW-LEVEL HELPERS
# ═══════════════════════════════════════════════════════════════════════

def _set_font(run, name=FONT, size=None, bold=None, italic=None):
    """Apply font formatting to a run."""
    run.font.name = name
    if size is not None:
        run.font.size = size
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = parse_xml(
            f'<w:rFonts {nsdecls("w")} '
            f'w:ascii="{name}" w:hAnsi="{name}" w:cs="{name}"/>'
        )
        rpr.insert(0, rf)
    else:
        for attr in ('w:ascii', 'w:hAnsi', 'w:cs'):
            rf.set(qn(attr), name)


def _add_run(para, text, size=Pt(10), bold=False, italic=False):
    """Add a formatted run to a paragraph."""
    run = para.add_run(text)
    _set_font(run, size=size, bold=bold, italic=italic)
    return run


def _set_spacing(para, after=0):
    pPr = para._element.get_or_add_pPr()
    sp = pPr.find(qn('w:spacing'))
    if sp is None:
        sp = parse_xml(f'<w:spacing {nsdecls("w")} w:after="{after}"/>')
        pPr.append(sp)
    else:
        sp.set(qn('w:after'), str(after))


def _set_border_bottom(para):
    pPr = para._element.get_or_add_pPr()
    bdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="auto"/>'
        f'</w:pBdr>'
    )
    pPr.insert(0, bdr)


def _setup_table(table):
    """Apply TableGrid style and full width to a table."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
    # Style
    if tblPr.find(qn('w:tblStyle')) is None:
        tblPr.insert(0, parse_xml(
            f'<w:tblStyle {nsdecls("w")} w:val="TableGrid"/>'
        ))
    # Width
    tw = tblPr.find(qn('w:tblW'))
    if tw is None:
        tblPr.append(parse_xml(
            f'<w:tblW {nsdecls("w")} w:w="{FULL_WIDTH}" w:type="dxa"/>'
        ))
    else:
        tw.set(qn('w:w'), str(FULL_WIDTH))
        tw.set(qn('w:type'), 'dxa')
    # Look
    if tblPr.find(qn('w:tblLook')) is None:
        tblPr.append(parse_xml(
            f'<w:tblLook {nsdecls("w")} w:val="04A0" '
            f'w:firstRow="1" w:lastRow="0" w:firstColumn="1" '
            f'w:lastColumn="0" w:noHBand="0" w:noVBand="1"/>'
        ))


def _set_cell_width(cell, width):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tw = tcPr.find(qn('w:tcW'))
    if tw is None:
        tcPr.append(parse_xml(
            f'<w:tcW {nsdecls("w")} w:w="{width}" w:type="dxa"/>'
        ))
    else:
        tw.set(qn('w:w'), str(width))
        tw.set(qn('w:type'), 'dxa')


def _merge_row(row):
    """Merge all cells in a row into the first cell."""
    cell = row.cells[0]
    if len(row.cells) > 1:
        cell.merge(row.cells[-1])
    return cell


def _set_grid_cols(table, widths):
    tbl = table._tbl
    grid = tbl.find(qn('w:tblGrid'))
    if grid is not None:
        tbl.remove(grid)
    xml = f'<w:tblGrid {nsdecls("w")}>'
    for w in widths:
        xml += f'<w:gridCol w:w="{w}"/>'
    xml += '</w:tblGrid>'
    grid = parse_xml(xml)
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is not None:
        tblPr.addnext(grid)
    else:
        tbl.insert(0, grid)


def _fmt_cell(cell, idx, text, size=Pt(10), bold=False, italic=False,
              align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold_prefix_colon=False):
    """Format paragraph `idx` in a cell (0 = existing, >0 = new)."""
    para = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
    if idx == 0:
        para.clear()
    para.alignment = align
    if text:
        if bold_prefix_colon and ":" in text:
            prefix, suffix = text.split(":", 1)
            _add_run(para, prefix + ":", size=size, bold=True)
            _add_run(para, suffix, size=size)
        else:
            _add_run(para, text, size=size, bold=bold, italic=italic)
    # Set paragraph-level default rPr
    pPr = para._element.get_or_add_pPr()
    if pPr.find(qn('w:rPr')) is None:
        rPr_xml = (
            f'<w:rPr {nsdecls("w")}>'
            f'<w:rFonts w:ascii="{FONT}" w:hAnsi="{FONT}" w:cs="{FONT}"/>'
        )
        if size:
            hp = int(size.pt * 2)
            rPr_xml += f'<w:sz w:val="{hp}"/><w:szCs w:val="{hp}"/>'
        rPr_xml += '</w:rPr>'
        pPr.append(parse_xml(rPr_xml))
    return para


def _fmt_label_value(cell, idx, label, value, size=Pt(10)):
    """Add a 'Label: value' line (label bold, value normal)."""
    para = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
    if idx == 0:
        para.clear()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _add_run(para, label, size=size, bold=True)
    _add_run(para, value, size=size)
    pPr = para._element.get_or_add_pPr()
    if pPr.find(qn('w:rPr')) is None:
        hp = int(size.pt * 2)
        pPr.append(parse_xml(
            f'<w:rPr {nsdecls("w")}>'
            f'<w:rFonts w:ascii="{FONT}" w:hAnsi="{FONT}" w:cs="{FONT}"/>'
            f'<w:sz w:val="{hp}"/><w:szCs w:val="{hp}"/>'
            f'</w:rPr>'
        ))
    return para


# ═══════════════════════════════════════════════════════════════════════
# DOCUMENT-LEVEL BUILDERS
# ═══════════════════════════════════════════════════════════════════════

def _add_section_header(doc, title):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_spacing(para, after=0)
    _set_border_bottom(para)
    _add_run(para, title, size=None, bold=True)


def _add_empty_para(doc):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_spacing(para, after=0)
    pPr = para._element.get_or_add_pPr()
    pPr.append(parse_xml(
        f'<w:rPr {nsdecls("w")}>'
        f'<w:rFonts w:ascii="{FONT}" w:hAnsi="{FONT}" w:cs="{FONT}"/>'
        f'</w:rPr>'
    ))


# ═══════════════════════════════════════════════════════════════════════
# SECTION GENERATORS  (each reads its portion of the JSON data)
# ═══════════════════════════════════════════════════════════════════════

def build_header(doc, data):
    """Build the header block (name, tagline, contact, linkedin)."""
    h = data.get("header", {})

    # Name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_spacing(p, 0)
    _add_run(p, h.get("full_name", ""), size=Pt(12))

    # Tagline (optional)
    tagline = h.get("tagline", "")
    if tagline:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_spacing(p, 0)
        _add_run(p, tagline, size=None)

    # Contact line
    parts = []
    if h.get("email"):
        parts.append(h["email"])
    if h.get("phone"):
        parts.append(h["phone"])
    if parts:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_spacing(p, 0)
        _add_run(p, " | ".join(parts), size=None)

    # Location (optional)
    if h.get("location"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_spacing(p, 0)
        _add_run(p, h["location"], size=None)

    # LinkedIn (optional)
    if h.get("linkedin"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_spacing(p, 0)
        _add_run(p, f'Linkedin: {h["linkedin"]}', size=None)

    _add_empty_para(doc)


def build_summary(doc, data):
    """Build the SUMMARY section."""
    s = data.get("summary", {})
    text = s.get("text", "")
    bold_words = s.get("bold_words", [])
    if not text:
        return

    _add_section_header(doc, "SUMMARY")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_spacing(p, 0)

    if bold_words:
        remaining = text
        for word in bold_words:
            if word in remaining:
                before, remaining = remaining.split(word, 1)
                if before:
                    _add_run(p, before)
                _add_run(p, word, bold=True)
        if remaining:
            _add_run(p, remaining)
    else:
        _add_run(p, text)

    _add_empty_para(doc)


def _build_two_col_entry(doc, entry, col_widths=None):
    """Build a generic 2-column table entry (education / work)."""
    if col_widths is None:
        col_widths = [TWO_COL_LEFT, TWO_COL_RIGHT]

    details = []
    # Collect label:value pairs
    for key, label in [("cgpa", "CGPA: "), ("results", "Results: "),
                       ("achievements", "Achievements: "),
                       ("coursework", "Relevant coursework: ")]:
        val = entry.get(key)
        if val:
            details.append((label, val))

    # Collect bullet descriptions
    for b in entry.get("bullets", []):
        details.append(b)  # plain string

    num_rows = 2 + len(details)
    table = doc.add_table(rows=num_rows, cols=2)
    _setup_table(table)
    _set_grid_cols(table, col_widths)

    # Row 0: Institution/Company + Location
    r0 = table.rows[0]
    _set_cell_width(r0.cells[0], col_widths[0])
    name_key = entry.get("institution") or entry.get("company") or entry.get("name", "")
    _fmt_cell(r0.cells[0], 0, name_key, bold=True)
    _set_cell_width(r0.cells[1], col_widths[1])
    _fmt_cell(r0.cells[1], 0, entry.get("location", ""), bold=True,
              align=WD_ALIGN_PARAGRAPH.RIGHT)

    # Row 1: Degree/Position + Date
    r1 = table.rows[1]
    _set_cell_width(r1.cells[0], col_widths[0])
    title_text = entry.get("degree") or entry.get("position") or ""
    _fmt_cell(r1.cells[0], 0, title_text, italic=True)
    _set_cell_width(r1.cells[1], col_widths[1])
    _fmt_cell(r1.cells[1], 0, entry.get("date_range", ""),
              align=WD_ALIGN_PARAGRAPH.RIGHT)

    # Detail rows (merged across 2 columns)
    for i, detail in enumerate(details):
        row = table.rows[2 + i]
        cell = _merge_row(row)
        _set_cell_width(cell, FULL_WIDTH)
        if isinstance(detail, tuple):
            _fmt_label_value(cell, 0, detail[0], detail[1])
        else:
            bullet_text = detail if detail.startswith("\u2022") else f"\u2022 {detail}"
            _fmt_cell(cell, 0, bullet_text, bold_prefix_colon=True)

    return table


def build_education(doc, data, key="education"):
    """Build EDUCATION section with any number of entries."""
    entries = data.get(key, [])
    if not entries:
        return
    _add_section_header(doc, data.get(f"{key}_header", "EDUCATION"))
    for i, entry in enumerate(entries):
        _build_two_col_entry(doc, entry)
        _add_empty_para(doc)


def build_work(doc, data, key="work"):
    """Build WORK EXPERIENCES section with any number of entries."""
    entries = data.get(key, [])
    if not entries:
        return
    header_text = data.get(f"{key}_header", "WORK EXPERIENCES")
    _add_section_header(doc, header_text)
    for i, entry in enumerate(entries):
        _build_two_col_entry(doc, entry)
        _add_empty_para(doc)


def build_projects(doc, data, key="projects"):
    """Build REPORT / ACADEMIC PROJECTS section."""
    section = data.get(key, {})
    entries = section.get("entries", [])
    if not entries:
        return
    header_text = section.get("header_text", "REPORT")
    _add_section_header(doc, header_text)
    for entry in entries:
        bullets = entry.get("bullets", [])
        num_rows = 1 + len(bullets)
        table = doc.add_table(rows=num_rows, cols=2)
        _setup_table(table)
        _set_grid_cols(table, [REPORT_LEFT, REPORT_RIGHT])

        r0 = table.rows[0]
        _set_cell_width(r0.cells[0], REPORT_LEFT)
        _fmt_cell(r0.cells[0], 0, entry.get("title", ""), bold=True)
        _set_cell_width(r0.cells[1], REPORT_RIGHT)
        _fmt_cell(r0.cells[1], 0, entry.get("date_range", ""), bold=True,
                  align=WD_ALIGN_PARAGRAPH.RIGHT)

        for i, bullet in enumerate(bullets):
            row = table.rows[1 + i]
            cell = _merge_row(row)
            _set_cell_width(cell, FULL_WIDTH)
            bullet_text = bullet if bullet.startswith("\u2022") else f"\u2022 {bullet}"
            _fmt_cell(cell, 0, bullet_text, bold_prefix_colon=True)

        _add_empty_para(doc)


def build_activities(doc, data, key="activities"):
    """Build ACTIVITIES & EXTRACURRICULAR section."""
    entries = data.get(key, [])
    if not entries:
        return
    _add_section_header(doc, data.get(f"{key}_header", "ACTIVITIES & EXTRACURRICULAR"))

    for entry in entries:
        bullets = entry.get("bullets", [])
        num_rows = 1 + max(1, len(bullets))
        
        table = doc.add_table(rows=num_rows, cols=2)
        _setup_table(table)
        _set_grid_cols(table, [REPORT_LEFT, REPORT_RIGHT])

        # Header row
        r0 = table.rows[0]
        _set_cell_width(r0.cells[0], REPORT_LEFT)
        _fmt_cell(r0.cells[0], 0, entry.get("name", ""), bold=True)
        _set_cell_width(r0.cells[1], REPORT_RIGHT)
        _fmt_cell(r0.cells[1], 0, entry.get("date_range", ""), bold=True,
                  align=WD_ALIGN_PARAGRAPH.RIGHT)

        # Bullet rows
        if not bullets:
            row = table.rows[1]
            cell = _merge_row(row)
            _set_cell_width(cell, FULL_WIDTH)
            _fmt_cell(cell, 0, "")
        else:
            for i, bullet in enumerate(bullets):
                row = table.rows[1 + i]
                cell = _merge_row(row)
                _set_cell_width(cell, FULL_WIDTH)
                bullet_text = bullet if bullet.startswith("\u2022") else f"\u2022 {bullet}"
                _fmt_cell(cell, 0, bullet_text, bold_prefix_colon=True)

        _add_empty_para(doc)


def build_certifications(doc, data, key="certifications"):
    """Build CERTIFICATIONS section (supports simple strings or detailed project-like dicts)."""
    section = data.get(key, {})
    entries = section.get("entries", [])
    if not entries:
        return
        
    if isinstance(entries[0], str):
        header_text = section.get("header_text", "CERTIFICATIONS AND PARTICIPATIONS")
        _add_section_header(doc, header_text)

        table = doc.add_table(rows=len(entries), cols=1)
        _setup_table(table)
        _set_grid_cols(table, [FULL_WIDTH])

        for i, text in enumerate(entries):
            row = table.rows[i]
            _set_cell_width(row.cells[0], FULL_WIDTH)
            cert = text if text.startswith("\u2022") else f"\u2022 {text}"
            _fmt_cell(row.cells[0], 0, cert)

        _add_empty_para(doc)
    else:
        # Use project builder for detailed certifications
        build_projects(doc, data, key)


def build_skills(doc, data, key="skills"):
    """Build ADDITIONAL INFORMATION section (skills + languages)."""
    section = data.get(key, {})
    categories = section.get("categories", [])
    if not categories:
        return
    header_text = section.get("header_text", "ADDITIONAL INFORMATION")
    _add_section_header(doc, header_text)

    table = doc.add_table(rows=1, cols=1)
    _setup_table(table)
    _set_grid_cols(table, [FULL_WIDTH])

    cell = table.rows[0].cells[0]
    _set_cell_width(cell, FULL_WIDTH)

    for i, cat in enumerate(categories):
        label = cat.get("label", "")
        value = cat.get("value", "")
        _fmt_label_value(cell, i, f"{label}: ", value)
        # Add empty separator between groups if there's a "separator" flag
        if cat.get("separator_after"):
            empty_p = cell.add_paragraph()
            empty_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    _add_empty_para(doc)


def build_references(doc, data, key="references"):
    """Build REFERENCES section (plain paragraphs, not tables)."""
    entries = data.get(key, [])
    if not entries:
        return
    _add_section_header(doc, data.get(f"{key}_header", "REFERENCES"))

    for i, ref in enumerate(entries):
        if i > 0:
            _add_empty_para(doc)  # separator between refs

        # Name (bold)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _set_spacing(p, 0)
        _add_run(p, ref.get("name", ""), size=None, bold=True)

        # Title
        if ref.get("title"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _set_spacing(p, 0)
            _add_run(p, ref["title"], size=None)

        # Phone
        if ref.get("phone"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _set_spacing(p, 0)
            _add_run(p, f'Office No: {ref["phone"]}', size=None)

        # Email
        if ref.get("email"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _set_spacing(p, 0)
            _add_run(p, f'Email: {ref["email"]}', size=None)


# ═══════════════════════════════════════════════════════════════════════
# ORCHESTRATOR
# ═══════════════════════════════════════════════════════════════════════

# Map section keys to builder functions
SECTION_BUILDERS = {
    "education":      build_education,
    "work":           build_work,
    "projects":       build_projects,
    "activities":     build_activities,
    "certifications": build_certifications,
    "skills":         build_skills,
    "references":     build_references,
}

# Default section order (used if not specified in JSON)
DEFAULT_ORDER = [
    "education", "work", "projects", "activities",
    "certifications", "skills", "references",
]


def generate_cv(data, output_path):
    """Generate a complete CV DOCX from structured data."""
    # Load template
    if TEMPLATE.exists():
        doc = Document(str(TEMPLATE))
    else:
        print(f"WARNING: Template not found at {TEMPLATE}, creating blank document")
        doc = Document()

    # Clear body (keep sectPr for page settings)
    body = doc.element.body
    sect_pr = body.find(qn('w:sectPr'))
    for child in list(body):
        if child.tag != qn('w:sectPr'):
            body.remove(child)

    # 1. Header (always first)
    build_header(doc, data)

    # 2. Summary (always second)
    build_summary(doc, data)

    # 3. Remaining sections in configured order
    order = data.get("sections_order", DEFAULT_ORDER)
    for section_key in order:
        base_key = section_key.split('-')[0]
        builder = SECTION_BUILDERS.get(base_key)
        if builder:
            builder(doc, data, section_key)

    # Move sectPr to end
    if sect_pr is not None:
        body.remove(sect_pr)
        body.append(sect_pr)

    # Save
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


# ═══════════════════════════════════════════════════════════════════════
# CLI
# ═══════════════════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(description="Universal CV Generator")
    parser.add_argument("--data", required=True, help="Path to JSON data file")
    parser.add_argument("--output", default=None,
                        help="Output filename (default: from JSON or cv_output.docx)")
    args = parser.parse_args()

    # Load JSON
    data_path = Path(args.data)
    if not data_path.is_absolute():
        # Try relative to data dir, then workspace
        candidates = [
            WORKSPACE / "data" / data_path,
            WORKSPACE / data_path,
            Path.cwd() / data_path,
        ]
        for c in candidates:
            if c.exists():
                data_path = c
                break

    if not data_path.exists():
        print(f"ERROR: Data file not found: {data_path}")
        sys.exit(1)

    with open(data_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    # Determine output path
    if args.output:
        out_name = args.output
    else:
        out_name = data.get("output_filename", "cv_output.docx")

    if not out_name.endswith(".docx"):
        out_name += ".docx"

    output_path = OUTPUT_DIR / out_name
    result = generate_cv(data, output_path)

    print(f"[OK] Generated: {result}")
    print(f"     File size: {result.stat().st_size:,} bytes")

    # Print section summary
    sections = []
    if data.get("education"):
        sections.append(f"Education ({len(data['education'])})")
    if data.get("work"):
        sections.append(f"Work ({len(data['work'])})")
    if data.get("projects", {}).get("entries"):
        sections.append(f"Projects ({len(data['projects']['entries'])})")
    if data.get("activities"):
        sections.append(f"Activities ({len(data['activities'])})")
    if data.get("certifications", {}).get("entries"):
        sections.append(f"Certifications ({len(data['certifications']['entries'])})")
    if data.get("skills", {}).get("categories"):
        sections.append(f"Skills ({len(data['skills']['categories'])} categories)")
    if data.get("references"):
        sections.append(f"References ({len(data['references'])})")

    print(f"   Sections: {', '.join(sections)}")


if __name__ == "__main__":
    main()
